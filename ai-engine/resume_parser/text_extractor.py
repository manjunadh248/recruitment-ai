"""
═══════════════════════════════════════════════════════════
RecruitAI — Text Extractor
═══════════════════════════════════════════════════════════
Extracts raw text from PDF and DOCX resume files using
PyMuPDF and python-docx respectively.
"""

import re
from pathlib import Path
from typing import Optional

from loguru import logger


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).
    Processes each page and concatenates the text.
    """
    import fitz  # PyMuPDF

    text_parts = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)
            logger.debug(f"Extracted page {page_num + 1}/{len(doc)}")
        doc.close()
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")

    raw = "\n".join(text_parts)
    return _clean_text(raw)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Includes paragraphs and table cell content.
    """
    from docx import Document

    text_parts = []
    try:
        doc = Document(file_path)

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")

    raw = "\n".join(text_parts)
    return _clean_text(raw)


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract text from a resume file based on its type.

    Args:
        file_path: Absolute path to the file
        file_type: 'pdf' or 'docx'

    Returns:
        Cleaned raw text string
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    file_type = file_type.lower().strip(".")

    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Use 'pdf' or 'docx'.")


def _clean_text(text: str) -> str:
    """Clean extracted text: normalize whitespace, remove artifacts."""
    # Replace multiple newlines with double newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Replace multiple spaces with single space
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Remove null bytes and other control characters (keep newlines/tabs)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    return text.strip()
